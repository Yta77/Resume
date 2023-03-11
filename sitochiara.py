import dash
import dash_bootstrap_components as dbc
from dash import dcc
from dash import html
from dash.dependencies import Input, Output
from dash import dash,html
import dash_bootstrap_components as dbc
from dash import dcc
import docx

app = dash.Dash(external_stylesheets=[dbc.themes.YETI])

#FILE READ FUNCTION

#TEXT FILE 0
def read_word_file(filepath):
    doc = docx.Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)



filepath = 'C:/Users/Workstation2/Desktop/Sitochiara/chisono.docx'

file_text = read_word_file(filepath)

#TEXT FILE 1

def read_word_file(filepath1):
    doc = docx.Document(filepath1)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

filepath1 = 'C:/Users/Workstation2/Desktop/Sitochiara/psicoforense.docx'

file_text1 = read_word_file(filepath1)

#TEXT FILE 2

def read_word_file(filepath2):
    doc = docx.Document(filepath2)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

filepath2 = 'C:/Users/Workstation2/Desktop/Sitochiara/Articolo1.docx'

file_text2 = read_word_file(filepath2)


#CONTACT US

contact_us = html.Div(
    [
              html.Div(
            [
                html.Span(className="fa fa-phone fa-lg"),
                html.Span("Phone: (+39) 388-1181074"),
            ],
            style={"margin-right": "20px"},
        ),
        html.Div(
            [
                html.Span(className="fa fa-map-marker fa-lg"),
                html.Span("Indirizzo: Via Lorenzo il Magnifico 110, Roma, 00162"),
            ],
            style={"margin-right": "20px"},
        ),
        html.Div(
            [
                html.Span(className="fa fa-envelope fa-lg"),
                html.Span("Email: info@chiarapaolapapi.it"),
            ],
        ),
    ]
)

#FOOTER

footer = html.Div([
    html.Footer(className="text-center text-white ",
                style={"background-color": "#21081a"},
                children=[
                    html.Div(className="container p-4"),
                    html.Div(className="text-center p-3",
                             style={"background-color": "rgba(0, 0, 0, 0.2)"},
                             children=[
                                 "© 2023 Copyright: ",
                                 html.A("Carlo Longo",
                                       # href="https://mdbootstrap.com/",
                                        className="text-white")
                             ])
                ])
])

#definizione carousel1

carousel = dbc.Carousel(
    items=[
        {"key": "1", "src": "/assets/roma3.jpg","img_style":{"max-height":"750px"}},
        {"key": "2", "src": "/assets/terapia.jpg","img_style":{"max-height":"750px"}},
        {"key": "3", "src": "/assets/bilancino.jpg","img_style":{"max-height":"750px"}}
    ],
    controls=True,
    indicators=True,
    interval=2000,
    ride="carousel"
)
#definizione carousel2

carousel2 = html.Div(
    [
        dbc.Carousel(
            id="carousel",
            items=[
                {"key": "1", "src": "/assets/chiara.jpeg",
                 "img_style": {"width": "800px", "height": "850px", "margin": "0 auto", "display": "block"},
                 "imgClassName": ""},
                {"key": "2", "src": "/assets/stanza.jpeg",
                 "img_style": {"width": "750px", "height": "850px", "margin": "0 auto", "display": "block"},

                 "imgClassName": ""},
            ],
            controls=True,
            indicators=True,
            interval=2000,
            ride="carousel",
        ),

    ]
)

#NAVIGATION BAR

navbar2 = dbc.NavbarSimple(
        children=[
                dbc.NavItem(dbc.NavLink("Chi sono", href="#start",external_link=True)),
                dbc.NavItem(dbc.NavLink("Contatti",href="#mid",external_link=True)),
                dbc.NavItem(dbc.NavLink("Dove Ricevo", href="#midstart",external_link=True)),
                dbc.NavItem(dbc.NavLink("Articoli", href="#midstart", external_link=True)),
                dbc.NavItem(dbc.NavLink("Servizi", href="#midend", external_link=True)),
                dbc.NavItem(dbc.NavLink("Consulenze Gestite", href="#end",external_link=True))],


    brand="Dott.ssa Chiara Paola Papi",
    brand_href="#",
    color="black",
    sticky='top',
    dark=True,

)

CONTENT_STYLE = {
   "margin-left": "3rem",
   "margin-right": "2rem",
   "padding": "2rem 1rem",
}


content = html.Div(
    [
        html.H1("La Psicologia Giuridica e Forense", id="start"),
        html.Br(),
        html.P("ddddddd \n aaaa",id='markdown0'),
        html.H1("Articolo", id="midstart"),
        html.P(id='markdown1', children=[file_text2]),
        carousel2,
        html.Br(),
        html.H1("Chi Sono", id="mid"),
        html.P(id='markdown2', children=[file_text]),
        html.Br(),
        html.H1("Servizi", id="midend"),
        html.P(dcc.Markdown('''
                                            * Consulenza Minori e Famiglia
                                            * Risarcimento del danno alla persona:
                                            * Danno da lutto
                                            * Danno da nascita indesiderata
                                            * Danno da wrongful life
                                            * Danno da menomazione fisica
                                            * Danno alla sfera sessuale
                                            * Danno da menomazione della capacità visiva
                                            * Danno estetico
                                            * Danno da mobbing
                                            * Danno da stalking
                                            * Danno da colpa professionale  
                                            * Danno da gaslhiting
                                            * Danno dei congiunti
                                            * Danno da carcerazione ingiusta
                                            * Danno da handicap
                                            * Idoneità per la ratificazione di attribuzione di sesso
                                            * Consulenza Mobbing e demansionamento
                                            * Valutazione dell'ammissibilità  o meno di una amministrazione di sostegno
                                            * Invalidità Civile
                                            * Capacità Testamentale
                                            * Capacità di intendere e volere

                                            '''))]+
        [html.H1("Contatti", id="end")]+
        [html.P(id='markdown3', children=(contact_us))]

                                     ),
#id="page-content", style=CONTENT_STYLE)

app.layout = html.Div([navbar2, carousel,content,footer])

if __name__ == '__main__':
    app.run_server(debug=True)